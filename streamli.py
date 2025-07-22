import streamlit as st
import torch
import torch.nn as nn
import torch.nn.functional as F
import numpy as np
import cv2
from PIL import Image
import io
import matplotlib.pyplot as plt
import os
import pandas as pd
from datetime import datetime
from transformers import AutoTokenizer, AutoModelForCausalLM
import json
import base64
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage, Table, TableStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from model import ResUNet50 

st.set_page_config(
    page_title="Brain Tumor Segmentation and Reporting App",
    page_icon="🧠",
    layout="wide"
)

class DoubleConv(nn.Module):
    def __init__(self, in_channels, out_channels):
        super(DoubleConv, self).__init__()
        self.conv = nn.Sequential(
            nn.Conv2d(in_channels, out_channels, kernel_size=3, padding=1),
            nn.BatchNorm2d(out_channels),
            nn.ReLU(inplace=True),
            nn.Conv2d(out_channels, out_channels, kernel_size=3, padding=1),
            nn.BatchNorm2d(out_channels),
            nn.ReLU(inplace=True)
        )

    def forward(self, x):
        return self.conv(x)

class UNet(nn.Module):
    def __init__(self, in_channels=3, out_channels=1, features=[64, 128, 256, 512]):
        super(UNet, self).__init__()
        self.downs = nn.ModuleList()
        self.ups = nn.ModuleList()
        self.pool = nn.MaxPool2d(kernel_size=2, stride=2)

        curr_channels = in_channels
        for feature in features:
            self.downs.append(DoubleConv(curr_channels, feature))
            curr_channels = feature

        self.bottleneck = DoubleConv(features[-1], features[-1] * 2)

        for feature in reversed(features):
            self.ups.append(nn.ConvTranspose2d(feature * 2, feature, kernel_size=2, stride=2))
            self.ups.append(DoubleConv(feature * 2, feature))

        self.final_conv = nn.Conv2d(features[0], out_channels, kernel_size=1)

    def forward(self, x):
        skip_connections = []
        for down in self.downs:
            x = down(x)
            skip_connections.append(x)
            x = self.pool(x)

        x = self.bottleneck(x)
        skip_connections = skip_connections[::-1]

        for idx in range(0, len(self.ups), 2):
            x = self.ups[idx](x)
            skip_connection = skip_connections[idx // 2]
            if x.shape != skip_connection.shape:
                x = F.interpolate(x, size=skip_connection.shape[2:])
            x = torch.cat((skip_connection, x), dim=1)
            x = self.ups[idx+1](x)

        return self.final_conv(x)

def preprocess_image(image_data, target_size=(256, 256)):
    try:
        image = Image.open(io.BytesIO(image_data))
        # Handle multi-page TIFF files by only using the first frame
        if hasattr(image, 'n_frames') and image.n_frames > 1:
            image.seek(0)  # Access first frame
        image = np.array(image)
    except Exception as e:
        st.error(f"Error opening image: {e}")
        return None, None
    
    original_image = image.copy()
    
    if len(image.shape) == 2:
        image = cv2.cvtColor(image, cv2.COLOR_GRAY2RGB)
    elif len(image.shape) == 3 and image.shape[2] > 3:
        image = image[:, :, :3]
    
    image_resized = cv2.resize(image, target_size)
    
    image_normalized = image_resized.astype(np.float32) / 255.0
    image_tensor = torch.tensor(image_normalized).permute(2, 0, 1).unsqueeze(0)
    
    return original_image, image_tensor
def predict(model, image_tensor, device):
    model.eval()
    with torch.no_grad():
        image_tensor = image_tensor.to(device)
        prediction = model(image_tensor)
        prediction = torch.sigmoid(prediction)
        prediction = (prediction > 0.5).float()
    return prediction

@st.cache_resource
def load_segmentation_model(model_file):
    model = ResUNet50(out_channels=1, pretrained=False)  # 🔁 Must match training config
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    
    model_bytes = io.BytesIO(model_file.getvalue())
    model.load_state_dict(torch.load(model_bytes, map_location=device))
    model.to(device)
    model.eval()
    return model, device

import requests
import streamlit as st

import streamlit as st
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM

@st.cache_resource
def load_report_generator_model(model_name):
    """
    Loads the model and tokenizer locally and returns a callable to generate reports.
    Cached using st.cache_resource to prevent reloading on every app rerun.
    """
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForCausalLM.from_pretrained(model_name)

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)

    def generate_report(prompt, max_length=256):
        inputs = tokenizer(prompt, return_tensors="pt").to(device)
        outputs = model.generate(**inputs, max_length=1024, do_sample=True)
        return tokenizer.decode(outputs[0], skip_special_tokens=True)

    return generate_report


    def generate_report(prompt):
        response = requests.post(
            f"https://api-inference.huggingface.co/models/{model_name}",
            headers=headers,
            json={"inputs": prompt}
        )
        
        if response.status_code == 200:
            return response.json()[0]['generated_text']
        else:
            st.error(f"Inference failed: {response.text}")
            return None

    return generate_report  # Only return the function



def overlay_mask(original_image, mask, alpha=0.5):
    if len(original_image.shape) == 2:
        original_image = cv2.cvtColor(original_image, cv2.COLOR_GRAY2RGB)
    
    if original_image.max() > 1.0:
        original_image = original_image.astype(np.float32) / 255.0
    
    overlay_img = original_image.copy()
    
    mask_resized = cv2.resize(mask, (original_image.shape[1], original_image.shape[0]), 
                              interpolation=cv2.INTER_NEAREST)
    
    colored_mask = np.zeros_like(overlay_img)
    if colored_mask.dtype != np.float32:
        colored_mask = colored_mask.astype(np.float32)
    
    colored_mask[mask_resized > 0] = [1.0, 0, 0]
    
    result = cv2.addWeighted(overlay_img, 1, colored_mask, alpha, 0)
    
    result = np.clip(result, 0, 1)
    
    return (result * 255).astype(np.uint8)

def prepare_for_display(image):
    if image.max() <= 1.0:
        image = (image * 255).astype(np.uint8)
    else:
        image = image.astype(np.uint8)
    return image

def generate_report(model, tokenizer, patient_info, segmentation_results):
    """Generate a medical report using the provided model and data"""
    try:
        prompt = f"""
Generate a comprehensive radiological report for a brain scan with the following details:

**Patient Demographics**
- Patient Name: {patient_info['name']}
- Patient ID: {patient_info['id']}
- Age: {patient_info['age']}
- Date of Birth: {patient_info['dob']}
- Sex: {patient_info['sex']}
- Contact Number: {patient_info['contact']}
- Email Address: {patient_info['email']}
- Address: {patient_info['address']}

**Medical History**
- Previous Medical Conditions: {patient_info['medical_history']}
- Family History: {patient_info['family_history']}
- Chief Complaint / Symptoms: {patient_info['chief_complaint']}
- Current Medications: {patient_info['current_medications']}
- Known Allergies: {patient_info['allergies']}
- Previous Imaging Studies: {patient_info['previous_imaging']}

**Clinical Information**
- Referring Physician: {patient_info['referring_physician']}
- Provisional Diagnosis: {patient_info['clinical_diagnosis']}
- Onset of Symptoms: {patient_info['onset_of_symptoms']}
- Symptom Progression: {patient_info['symptom_progression']}
- Neurological Symptoms: {patient_info['neurological_symptoms']}
- Treatment History: {patient_info['treatment_history']}
- Additional Clinical Notes: {patient_info['additional_notes']}

**Scan Findings**
- Affected Brain Area (%): {st.session_state['segmentation_results']['affected_percentage']:.2f}%
- Location of Lesion: {st.session_state['segmentation_results']['location']}
- Number of Images Analyzed: {st.session_state['segmentation_results']['num_images']}
- Date of Study: {datetime.now().strftime('%Y-%m-%d')}

Based on the above information, provide a detailed radiological assessment organized as follows:

1. **Clinical History:**  
   Summarize the patient’s presenting symptoms, relevant past history, and any risk factors.

2. **Technique:**  
   Describe the imaging modality, sequences used, and any contrast administration.

3. **Findings:**  
   Detail the lesion’s size, percentage of involvement, signal/intensity characteristics, mass effect, edema, and any other notable observations.

4. **Impression with Differential Diagnosis:**  
   Offer one or two most likely diagnoses and list differential considerations.

5. **Recommendations & Follow-Up:**  
   Suggest further imaging or clinical actions, with a recommended timeframe (e.g., “Follow-up MRI in 3–6 months”).

6. **Additional Comments (if any):**  
   Include any miscellaneous remarks or questions for the referring physician.
"""
        inputs = tokenizer(prompt, return_tensors="pt")
        with torch.no_grad():
            output = model.generate(
                inputs["input_ids"],
                max_length=1024,
                num_return_sequences=1,
                temperature=0.7
            )
            
        report = tokenizer.decode(output[0], skip_special_tokens=True)
        return report
    except Exception as e:
        st.error(f"Error generating report: {e}")
        return "Failed to generate report due to an error."

def determine_tumor_location(mask, original_shape):
    """Determine the approximate location of the tumor in the brain"""
    # This is a simplistic approach - in a real application, you would use a more sophisticated method
    h, w = mask.shape
    
    # Find center of mass of the tumor
    y_indices, x_indices = np.where(mask > 0)
    
    if len(y_indices) == 0 or len(x_indices) == 0:
        return "No tumor detected"
    
    center_y = np.mean(y_indices) / h
    center_x = np.mean(x_indices) / w
    
    # Determine region based on center of mass
    region = []
    
    # Vertical position
    if center_y < 0.33:
        region.append("Superior")
    elif center_y > 0.66:
        region.append("Inferior")
    else:
        region.append("Mid")
    
    # Horizontal position
    if center_x < 0.33:
        region.append("Left")
    elif center_x > 0.66:
        region.append("Right")
    else:
        region.append("Central")
    
    return " ".join(region) + " region"

def calculate_tumor_metrics(mask):
    """Calculate additional metrics for the tumor"""
    if mask.sum() == 0:
        return {
            "volume_estimate": 0,
            "max_diameter": 0,
            "shape_irregularity": 0
        }
    
    # Simple volume estimate (pixel count as proxy)
    volume = mask.sum()
    
    # Estimate maximum diameter
    y_indices, x_indices = np.where(mask > 0)
    if len(y_indices) == 0:
        return {"volume_estimate": 0, "max_diameter": 0, "shape_irregularity": 0}
    
    # Approximate bounding box
    min_y, max_y = np.min(y_indices), np.max(y_indices)
    min_x, max_x = np.min(x_indices), np.max(x_indices)
    height = max_y - min_y
    width = max_x - min_x
    max_diameter = max(height, width)
    
    # Shape irregularity (ratio of perimeter squared to area, normalized)
    # This is a simple approximation
    contours, _ = cv2.findContours((mask * 255).astype(np.uint8), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if contours:
        perimeter = cv2.arcLength(contours[0], True)
        area = cv2.contourArea(contours[0])
        if area > 0:
            irregularity = (perimeter ** 2) / (4 * np.pi * area)
        else:
            irregularity = 0
    else:
        irregularity = 0
    
    return {
        "volume_estimate": volume,
        "max_diameter": max_diameter,
        "shape_irregularity": irregularity
    }

def create_pdf_report(patient_info, segmentation_results, report_text, images=None):
    """Create a PDF report with proper medical formatting and enhanced patient details"""
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=1,  # Center alignment
        spaceAfter=12
    )
    
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=6
    )
    
    subheading_style = ParagraphStyle(
        'SubHeading',
        parent=styles['Heading3'],
        fontSize=11,
        spaceAfter=3
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    
    # Build the document content
    content = []
    
    # Header/Title
    content.append(Paragraph("RADIOLOGICAL REPORT", title_style))
    content.append(Paragraph("Brain MRI Analysis", title_style))
    content.append(Spacer(1, 0.2*inch))
    
    # Institution info
    content.append(Paragraph("Medical Imaging Center", heading_style))
    content.append(Paragraph("123 Healthcare Avenue, Medical City", normal_style))
    content.append(Paragraph("Phone: (555) 123-4567", normal_style))
    content.append(Spacer(1, 0.2*inch))
    
    # Basic Patient Info
    content.append(Paragraph("PATIENT INFORMATION", heading_style))
    
    basic_data = [
        ["Patient Name:", patient_info.get('name', 'N/A'), "Patient ID:", patient_info.get('id', 'N/A')],
        ["Date of Birth:", patient_info.get('dob', 'N/A').strftime('%Y-%m-%d') if isinstance(patient_info.get('dob'), datetime.date) else 'N/A', 
         "Age:", str(patient_info.get('age', 'N/A'))],
        ["Sex:", patient_info.get('sex', 'N/A'), "Contact:", patient_info.get('contact', 'N/A')],
        ["Address:", patient_info.get('address', 'N/A'), "Email:", patient_info.get('email', 'N/A')]
    ]
    
    basic_table = Table(basic_data, colWidths=[1.2*inch, 1.8*inch, 1.2*inch, 1.8*inch])
    basic_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('BACKGROUND', (2, 0), (2, -1), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    content.append(basic_table)
    content.append(Spacer(1, 0.2*inch))
    
    # Clinical Information
    content.append(Paragraph("CLINICAL INFORMATION", heading_style))
    
    # Chief complaint and diagnosis
    clinical_data = [
        ["Chief Complaint:", patient_info.get('chief_complaint', 'N/A')],
        ["Provisional Diagnosis:", patient_info.get('clinical_diagnosis', 'N/A')],
        ["Referring Physician:", patient_info.get('referring_physician', 'N/A')],
        ["Onset of Symptoms:", patient_info.get('onset_of_symptoms', 'N/A')],
        ["Symptom Progression:", patient_info.get('symptom_progression', 'N/A')],
        ["Neurological Symptoms:", patient_info.get('neurological_symptoms', 'N/A')]
    ]
    
    clinical_table = Table(clinical_data, colWidths=[2*inch, 4*inch])
    clinical_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    content.append(clinical_table)
    content.append(Spacer(1, 0.1*inch))
    
    # Medical History
    content.append(Paragraph("Medical History", subheading_style))
    content.append(Paragraph(patient_info.get('medical_history', 'No medical history provided.'), normal_style))
    
    content.append(Paragraph("Family History", subheading_style))
    content.append(Paragraph(patient_info.get('family_history', 'No family history provided.'), normal_style))
    
    content.append(Paragraph("Current Medications", subheading_style))
    content.append(Paragraph(patient_info.get('current_medications', 'No current medications noted.'), normal_style))
    
    content.append(Paragraph("Allergies", subheading_style))
    content.append(Paragraph(patient_info.get('allergies', 'No known allergies.'), normal_style))
    
    content.append(Paragraph("Previous Imaging", subheading_style))
    content.append(Paragraph(patient_info.get('previous_imaging', 'No previous imaging noted.'), normal_style))
    
    content.append(Paragraph("Treatment History", subheading_style))
    content.append(Paragraph(patient_info.get('treatment_history', 'No treatment history provided.'), normal_style))
    
    content.append(Paragraph("Additional Notes", subheading_style))
    content.append(Paragraph(patient_info.get('additional_notes', 'No additional notes.'), normal_style))
    
    content.append(Spacer(1, 0.2*inch))
    
    # Format the report text - split by sections
    content.append(Paragraph("IMAGING REPORT", heading_style))
    if "FINDINGS" in report_text:
        parts = report_text.split("FINDINGS")
        if len(parts) > 1:
            preamble = parts[0].strip()
            rest = "FINDINGS" + parts[1]
            content.append(Paragraph(preamble, normal_style))
            
            # Process the findings and other sections
            sections = ["FINDINGS", "IMPRESSION", "RECOMMENDATIONS"]
            current_text = rest
            
            for i, section in enumerate(sections):
                if section in current_text:
                    parts = current_text.split(section)
                    if len(parts) > 1:
                        if i > 0:  # Don't add empty content for the first split
                            content.append(Paragraph(parts[0].strip(), normal_style))
                        content.append(Paragraph(section, heading_style))
                        current_text = section + parts[1]
                    else:
                        content.append(Paragraph(current_text, normal_style))
                        break
            
            # Add the remaining text
            content.append(Paragraph(current_text.replace(sections[-1], "").strip(), normal_style))
        else:
            # If we can't split the text, just add it all
            content.append(Paragraph(report_text, normal_style))
    else:
        # If we can't identify sections, just add the full text
        content.append(Paragraph(report_text, normal_style))
    
    # Add images if available
    if images:
        content.append(Paragraph("IMAGES", heading_style))
        # Add up to 4 images (original and segmented)
        for i, img_data in enumerate(images[:2]):  # Limit to 2 pairs
            # Create a 2-column table for each image pair
            img_table_data = []
            
            # Convert numpy arrays to reportlab images
            original_img_path = f"/tmp/original_{i}.png"
            cv2.imwrite(original_img_path, cv2.cvtColor(img_data['original'], cv2.COLOR_RGB2BGR))
            
            overlay_img_path = f"/tmp/overlay_{i}.png"
            cv2.imwrite(overlay_img_path, cv2.cvtColor(img_data['overlay'], cv2.COLOR_RGB2BGR))
            
            img_original = ReportLabImage(original_img_path, width=2*inch, height=2*inch)
            img_overlay = ReportLabImage(overlay_img_path, width=2*inch, height=2*inch)
            
            img_table_data.append([img_original, img_overlay])
            img_table_data.append([
                Paragraph(f"Original - {img_data['filename']}", normal_style),
                Paragraph(f"Segmentation - {img_data['filename']}", normal_style)
            ])
            
            img_table = Table(img_table_data, colWidths=[2.5*inch, 2.5*inch])
            img_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            content.append(img_table)
            content.append(Spacer(1, 0.2*inch))
    
    # Add segmentation results summary
    content.append(Paragraph("QUANTITATIVE ANALYSIS", heading_style))
    
    quant_data = [
        ["Affected Area:", f"{segmentation_results['affected_percentage']:.2f}%"],
        ["Location:", segmentation_results['location']],
        ["Volume Estimate:", f"{segmentation_results.get('volume_estimate', 'N/A')} pixels"],
        ["Maximum Diameter:", f"{segmentation_results.get('max_diameter', 'N/A')} pixels"],
        ["Shape Irregularity:", f"{segmentation_results.get('shape_irregularity', 'N/A'):.2f}"]
    ]
    
    quant_table = Table(quant_data, colWidths=[2*inch, 3*inch])
    quant_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    content.append(quant_table)
    content.append(Spacer(1, 0.2*inch))
    
    # Signature
    content.append(Paragraph("INTERPRETING RADIOLOGIST", heading_style))
    content.append(Paragraph("______________________________", normal_style))
    content.append(Paragraph("[Radiologist Name], MD", normal_style))
    content.append(Paragraph("Board Certified Radiologist", normal_style))
    content.append(Paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
    
    # Footer
    def add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        footer_text = f"Generated by Brain Tumor Segmentation & Reporting System - {datetime.now().strftime('%Y-%m-%d')}"
        canvas.drawCentredString(letter[0]/2, 0.5*inch, footer_text)
        canvas.drawCentredString(letter[0]/2, 0.25*inch, "Page %d" % canvas.getPageNumber())
        canvas.restoreState()
    
    # Build the document
    doc.build(content, onFirstPage=add_footer, onLaterPages=add_footer)
    
    buffer.seek(0)
    return buffer
def create_docx_report(patient_info, segmentation_results, report_text, images=None):
    """Create a DOCX report with proper medical formatting"""
    doc = Document()
    
    # Add header/title
    title = doc.add_heading('RADIOLOGICAL REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('Brain MRI Analysis', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add institution info
    doc.add_paragraph('Medical Imaging Center')
    doc.add_paragraph('123 Healthcare Avenue, Medical City')
    doc.add_paragraph('Phone: (555) 123-4567')
    doc.add_paragraph('')
    
    # Add patient information table
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    
    # Row 1
    row = table.rows[0]
    row.cells[0].text = 'Patient Name:'
    row.cells[1].text = patient_info['name']
    row.cells[2].text = 'Patient ID:'
    row.cells[3].text = patient_info['id']
    
    # Row 2
    row = table.rows[1]
    row.cells[0].text = 'Age:'
    row.cells[1].text = str(patient_info['age'])
    row.cells[2].text = 'Sex:'
    row.cells[3].text = patient_info['sex']
    
    # Row 3
    row = table.rows[2]
    row.cells[0].text = 'Study Date:'
    row.cells[1].text = datetime.now().strftime('%Y-%m-%d')
    row.cells[2].text = 'Report Date:'
    row.cells[3].text = datetime.now().strftime('%Y-%m-%d')
    
    doc.add_paragraph('')
    
    # Clinical History
    doc.add_heading('CLINICAL HISTORY', level=2)
    doc.add_paragraph(patient_info['medical_history'] if patient_info['medical_history'] else "No clinical history provided.")
    
    # Format the report text - split by sections
    if "FINDINGS" in report_text:
        parts = report_text.split("FINDINGS")
        if len(parts) > 1:
            preamble = parts[0].strip()
            rest = "FINDINGS" + parts[1]
            doc.add_paragraph(preamble)
            
            # Process the findings and other sections
            sections = ["FINDINGS", "IMPRESSION", "RECOMMENDATIONS"]
            current_text = rest
            
            for i, section in enumerate(sections):
                if section in current_text:
                    parts = current_text.split(section)
                    if len(parts) > 1:
                        if i > 0:  # Don't add empty content for the first split
                            doc.add_paragraph(parts[0].strip())
                        doc.add_heading(section, level=2)
                        current_text = section + parts[1]
                    else:
                        doc.add_paragraph(current_text)
                        break
            
            # Add the remaining text
            doc.add_paragraph(current_text.replace(sections[-1], "").strip())
        else:
            # If we can't split the text, just add it all
            doc.add_paragraph(report_text)
    else:
        # If we can't identify sections, just add the full text
        doc.add_paragraph(report_text)
    
    # Add images if available
    if images:
        doc.add_heading('IMAGES', level=2)
        # Add up to 4 images (original and segmented)
        for i, img_data in enumerate(images[:2]):  # Limit to 2 pairs
            # Save images temporarily
            original_img_path = f"/tmp/original_{i}.png"
            overlay_img_path = f"/tmp/overlay_{i}.png"
            
            cv2.imwrite(original_img_path, cv2.cvtColor(img_data['original'], cv2.COLOR_RGB2BGR))
            cv2.imwrite(overlay_img_path, cv2.cvtColor(img_data['overlay'], cv2.COLOR_RGB2BGR))
            
            # Add images
            table = doc.add_table(rows=2, cols=2)
            
            cell = table.cell(0, 0)
            cell.add_paragraph().add_run().add_picture(original_img_path, width=Inches(2.5))
            
            cell = table.cell(0, 1)
            cell.add_paragraph().add_run().add_picture(overlay_img_path, width=Inches(2.5))
            
            cell = table.cell(1, 0)
            cell.text = f"Original - {img_data['filename']}"
            
            cell = table.cell(1, 1)
            cell.text = f"Segmentation - {img_data['filename']}"
            
            doc.add_paragraph('')
    
    # Add segmentation results summary
    doc.add_heading('QUANTITATIVE ANALYSIS', level=2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # Add data to table
    data_rows = [
        ["Affected Area:", f"{segmentation_results['affected_percentage']:.2f}%"],
        ["Location:", segmentation_results['location']],
        ["Volume Estimate:", f"{segmentation_results.get('volume_estimate', 'N/A')} pixels"],
        ["Maximum Diameter:", f"{segmentation_results.get('max_diameter', 'N/A')} pixels"],
        ["Shape Irregularity:", f"{segmentation_results.get('shape_irregularity', 'N/A'):.2f}"]
    ]
    
    for i, row_data in enumerate(data_rows):
        row = table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
    
    doc.add_paragraph('')
    
    # Signature
    doc.add_heading('INTERPRETING RADIOLOGIST', level=2)
    doc.add_paragraph('______________________________')
    doc.add_paragraph('[Radiologist Name], MD')
    doc.add_paragraph('Board Certified Radiologist')
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Footer
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = f"Generated by Brain Tumor Segmentation & Reporting System - {datetime.now().strftime('%Y-%m-%d')}"
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_html_report(patient_info, segmentation_results, report_text, images=None):
    """Create an HTML report with proper medical formatting"""
    
    # Convert any image data to base64 for embedding
    image_html = ""
    if images:
        image_html = "<h2>IMAGES</h2><div style='display: flex; flex-wrap: wrap;'>"
        for i, img_data in enumerate(images[:2]):  # Limit to 2 pairs
            # Convert images to base64
            _, original_buffer = cv2.imencode('.png', cv2.cvtColor(img_data['original'], cv2.COLOR_RGB2BGR))
            original_base64 = base64.b64encode(original_buffer).decode('utf-8')
            
            _, overlay_buffer = cv2.imencode('.png', cv2.cvtColor(img_data['overlay'], cv2.COLOR_RGB2BGR))
            overlay_base64 = base64.b64encode(overlay_buffer).decode('utf-8')
            
            image_html += f"""
            <div style='margin-right: 20px; margin-bottom: 20px;'>
                <div style='display: flex;'>
                    <div style='margin-right: 10px;'>
                        <img src='data:image/png;base64,{original_base64}' style='width: 300px;'>
                        <p>Original - {img_data['filename']}</p>
                    </div>
                    <div>
                        <img src='data:image/png;base64,{overlay_base64}' style='width: 300px;'>
                        <p>Segmentation - {img_data['filename']}</p>
                    </div>
                </div>
            </div>
            """
        image_html += "</div>"
    
    # Format the report text with proper sections
    formatted_report = report_text
    if "FINDINGS" in report_text:
        sections = ["FINDINGS", "IMPRESSION", "RECOMMENDATIONS"]
        for section in sections:
            if section in formatted_report:
                formatted_report = formatted_report.replace(section, f"<h2>{section}</h2>")
    
    # Build the complete HTML document
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Radiological Report - {patient_info['name']}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
            .header {{ text-align: center; margin-bottom: 20px; }}
            .institution {{ margin-bottom: 20px; }}
            table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; }}
            th {{ background-color: #f2f2f2; text-align: left; width: 30%; }}
            h1 {{ color: #2c3e50; }}
            h2 {{ color: #3498db; margin-top: 20px; }}
            .footer {{ margin-top: 50px; text-align: center; font-size: 0.8em; color: #7f8c8d; }}
            .signature {{ margin-top: 30px; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>RADIOLOGICAL REPORT</h1>
            <h2>Brain MRI Analysis</h2>
        </div>
        
        <div class="institution">
            <p>Medical Imaging Center<br>
            123 Healthcare Avenue, Medical City<br>
            Phone: (555) 123-4567</p>
        </div>
        
        <table>
            <tr>
                <th>Patient Name:</th>
                <td>{patient_info['name']}</td>
                <th>Patient ID:</th>
                <td>{patient_info['id']}</td>
            </tr>
            <tr>
                <th>Age:</th>
                <td>{patient_info['age']}</td>
                <th>Sex:</th>
                <td>{patient_info['sex']}</td>
            </tr>
            <tr>
                <th>Study Date:</th>
                <td>{datetime.now().strftime('%Y-%m-%d')}</td>
                <th>Report Date:</th>
                <td>{datetime.now().strftime('%Y-%m-%d')}</td>
            </tr>
        </table>
        
        <h2>CLINICAL HISTORY</h2>
        <p>{patient_info['medical_history'] if patient_info['medical_history'] else "No clinical history provided."}</p>
        
        <div class="report-content">
            {formatted_report}
        </div>
        
        {image_html}
        
        <h2>QUANTITATIVE ANALYSIS</h2>
        <table>
            <tr>
                <th>Affected Area:</th>
                <td>{segmentation_results['affected_percentage']:.2f}%</td>
            </tr>
            <tr>
                <th>Location:</th>
                <td>{segmentation_results['location']}</td>
            </tr>
            <tr>
                <th>Volume Estimate:</th>
                <td>{segmentation_results.get('volume_estimate', 'N/A')} pixels</td>
            </tr>
            <tr>
                <th>Maximum Diameter:</th>
                <td>{segmentation_results.get('max_diameter', 'N/A')} pixels</td>
            </tr>
            <tr>
                <th>Shape Irregularity:</th>
                <td>{segmentation_results.get('shape_irregularity', 'N/A'):.2f}</td>
            </tr>
        </table>
        
        <div class="signature">
            <h2>INTERPRETING RADIOLOGIST</h2>
            <p>______________________________</p>
            <p>[Radiologist Name], MD</p>
            <p>Board Certified Radiologist</p>
            <p>Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        </div>
        
        <div class="footer">
            <p>Generated by Brain Tumor Segmentation & Reporting System - {datetime.now().strftime('%Y-%m-%d')}</p>
        </div>
    </body>
    </html>
    """
    
    return html

def main():
    st.title("🧠 Brain Tumor Segmentation and Reporting System")
    
    # Sidebar
    st.sidebar.title("Navigation")
    app_mode = st.sidebar.selectbox("Choose the app mode", 
        ["Home", "Single Image Analysis", "Batch Processing", "Generate Report"])
    uploaded_file = None  # ✅ Declare this outside any condition

    # Home page
    if app_mode == "Home":
        st.markdown("""
        ## Welcome to the Brain Tumor Segmentation and Reporting System
        
        This application allows medical professionals to:
        
        - Analyze brain MRI images for tumor detection
        - Segment tumors from MRI scans
        - Generate comprehensive medical reports
        - Process multiple images in batch mode
        - Export findings in various formats (PDF, DOCX, HTML)
        
        ### Getting Started
        1. Select a mode from the sidebar
        2. Upload your MRI scan(s)
        3. Review the segmentation results
        4. Generate a customized report
        
        ### About
        This system uses a U-Net deep learning model for tumor segmentation, with additional quantitative analysis capabilities.
        """)
        
        st.info("Please select a mode from the sidebar to begin.")
    
    # Single image analysis mode
    elif app_mode == "Single Image Analysis":
        st.header("Single MRI Image Analysis")
        uploaded_file = None

        # File uploader
        uploaded_file = st.file_uploader("Upload an MRI scan", type=["jpg", "jpeg", "png", "tif", "tiff"])
        
    if uploaded_file is not None:
            # Load segmentation model
            model_file = st.sidebar.file_uploader("Upload model file (optional)", type=["pth"])
            
            if model_file is None:
                st.warning("No model file uploaded. Please upload a trained model file to perform segmentation.")
                return
            
            # Process image

            model, device = load_segmentation_model(model_file)
            
            # Read image
            image_data = uploaded_file.getvalue()
            original_image, image_tensor = preprocess_image(image_data)
            
            if original_image is None or image_tensor is None:
                st.error("Failed to process the uploaded image. Please try a different image.")
                return
            
            # Show original image
            st.subheader("Original MRI Scan")
            st.image(prepare_for_display(original_image), caption="Original MRI Scan", use_column_width=True)
            
            # Perform segmentation
            st.subheader("Segmentation Results")
            with st.spinner("Performing segmentation..."):
                mask_pred = predict(model, image_tensor, device)
                mask_numpy = mask_pred[0, 0].cpu().numpy()
                
                # Create overlay image
                overlay = overlay_mask(original_image, mask_numpy)
                st.image(overlay, caption="Tumor Segmentation Overlay", use_column_width=True)
                
                # Calculate metrics
                tumor_pixels = np.sum(mask_numpy > 0.5)
                total_pixels = mask_numpy.size
                affected_percentage = (tumor_pixels / total_pixels) * 100
                
                location = determine_tumor_location(mask_numpy, original_image.shape)
                additional_metrics = calculate_tumor_metrics(mask_numpy)
                
                # Display metrics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Affected Area", f"{affected_percentage:.2f}%")
                with col2:
                    st.metric("Location", location)
                with col3:
                    st.metric("Volume Estimate", f"{additional_metrics['volume_estimate']}")
                
                # Save results in session state for report generation
                st.session_state['segmentation_results'] = {
                    'affected_percentage': affected_percentage,
                    'location': location,
                    'num_images': 1,
                    'volume_estimate': additional_metrics['volume_estimate'],
                    'max_diameter': additional_metrics['max_diameter'],
                    'shape_irregularity': additional_metrics['shape_irregularity']
                }
                
                st.session_state['image_data'] = [{
                    'original': original_image,
                    'overlay': overlay,
                    'filename': uploaded_file.name
                }]
                
                st.success("Segmentation completed successfully!")
                
                # Option to generate report
                if st.button("Proceed to Report Generation"):
                    st.session_state['app_mode'] = "Generate Report"
                    st.rerun()
    
    # Batch processing mode
    elif app_mode == "Batch Processing":
        st.header("Batch MRI Image Processing")
        
        # File uploader for multiple files
        uploaded_files = st.file_uploader("Upload multiple MRI scans", type=["jpg", "jpeg", "png", "tif", "tiff"], accept_multiple_files=True)
        
        if uploaded_files:
            # Load segmentation model
            model_file = st.sidebar.file_uploader("Upload model file (optional)", type=["pth"])
            
            if model_file is None:
                st.warning("No model file uploaded. Please upload a trained model file to perform segmentation.")
                return
            
            # Process images
            model, device = load_segmentation_model(model_file)
            
            # Initialize results storage
            results = []
            image_data = []
            total_affected = 0
            
            # Process each image
            for i, file in enumerate(uploaded_files):
                # Read image
                image_bytes = file.getvalue()
                original_image, image_tensor = preprocess_image(image_bytes)
                
                if original_image is None or image_tensor is None:
                    st.error(f"Failed to process image {file.name}. Skipping.")
                    continue
                
                # Perform segmentation
                mask_pred = predict(model, image_tensor, device)
                mask_numpy = mask_pred[0, 0].cpu().numpy()
                
                # Create overlay image
                overlay = overlay_mask(original_image, mask_numpy)
                
                # Calculate metrics
                tumor_pixels = np.sum(mask_numpy > 0.5)
                total_pixels = mask_numpy.size
                affected_percentage = (tumor_pixels / total_pixels) * 100
                
                location = determine_tumor_location(mask_numpy, original_image.shape)
                additional_metrics = calculate_tumor_metrics(mask_numpy)
                
                # Store results
                results.append({
                    'filename': file.name,
                    'affected_percentage': affected_percentage,
                    'location': location,
                    'volume_estimate': additional_metrics['volume_estimate'],
                    'max_diameter': additional_metrics['max_diameter'],
                    'shape_irregularity': additional_metrics['shape_irregularity']
                })
                
                image_data.append({
                    'original': original_image,
                    'overlay': overlay,
                    'filename': file.name
                })
                
                total_affected += affected_percentage
            
            # Display results in table
            st.subheader("Batch Processing Results")
            
            if results:
                results_df = pd.DataFrame(results)
                st.dataframe(results_df)
                
                # Display average metrics
                avg_affected = total_affected / len(results)
                st.metric("Average Affected Area", f"{avg_affected:.2f}%")
                
                # Save results in session state for report generation
                st.session_state['segmentation_results'] = {
                    'affected_percentage': avg_affected,
                    'location': "Multiple regions",
                    'num_images': len(results),
                    'volume_estimate': sum(r['volume_estimate'] for r in results) / len(results),
                    'max_diameter': max(r['max_diameter'] for r in results),
                    'shape_irregularity': sum(r['shape_irregularity'] for r in results) / len(results)
                }
                
                st.session_state['image_data'] = image_data
                
                # Option to view individual results
                with st.expander("View Individual Image Results"):
                    for i, (img_data, result) in enumerate(zip(image_data, results)):
                        st.subheader(f"Image {i+1}: {result['filename']}")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.image(prepare_for_display(img_data['original']), caption="Original MRI Scan", use_column_width=True)
                        with col2:
                            st.image(img_data['overlay'], caption="Tumor Segmentation Overlay", use_column_width=True)
                        
                        st.metric("Affected Area", f"{result['affected_percentage']:.2f}%")
                        st.metric("Location", result['location'])
                        st.metric("Volume Estimate", f"{result['volume_estimate']}")
                
                st.success(f"Successfully processed {len(results)} images!")
                
                # Option to generate report
                if st.button("Proceed to Report Generation"):
                    st.session_state['app_mode'] = "Generate Report"
                    st.rerun()
            else:
                st.error("No images were successfully processed.")
    
    # Report generation mode
    elif app_mode == "Generate Report":
        st.header("Generate Medical Report")
        
        # Check if results are available
        if 'segmentation_results' not in st.session_state:
            st.warning("No analysis results available. Please perform segmentation first.")
            return
        
        # Patient information form
       # Replace the existing patient information form section with this expanded version:

st.subheader("Patient Information")
with st.expander("Patient Demographics", expanded=True):
    col1, col2, col3 = st.columns(3)
    with col1:
        patient_name = st.text_input("Patient Name", "John Doe")
        patient_age = st.number_input("Age", 0, 120, 45)
        patient_dob = st.date_input("Date of Birth")
    
    with col2:
        patient_id = st.text_input("Patient ID", "P12345")
        patient_sex = st.selectbox("Sex", ["Male", "Female", "Other"])
        patient_contact = st.text_input("Contact Number")
    
    with col3:
        patient_email = st.text_input("Email Address")
        patient_address = st.text_area("Address", height=100)

with st.expander("Medical History", expanded=True):
    col1, col2 = st.columns(2)
    
    with col1:
        medical_history = st.text_area("Previous Medical Conditions", 
                                      height=100, 
                                      placeholder="List any significant medical conditions...")
        
        family_history = st.text_area("Family History", 
                                    height=100,
                                    placeholder="Relevant family medical history...")
        
        chief_complaint = st.text_area("Chief Complaint/Symptoms", 
                                     height=100,
                                     placeholder="Main reason for current examination...")
    
    with col2:
        current_medications = st.text_area("Current Medications", 
                                         height=100,
                                         placeholder="List current medications and dosages...")
        
        allergies = st.text_area("Known Allergies", 
                               height=68,
                               placeholder="Medications, contrast agents, etc...")
        
        previous_imaging = st.text_area("Previous Imaging Studies", 
                                      height=100,
                                      placeholder="Prior scans, dates, findings...")

with st.expander("Clinical Information", expanded=True):
    col1, col2 = st.columns(2)
    
    with col1:
        referring_physician = st.text_input("Referring Physician")
        clinical_diagnosis = st.text_input("Provisional Diagnosis")
        
        onset_of_symptoms = st.selectbox("Onset of Symptoms", 
                                       ["Recent (< 1 week)", "Subacute (1-4 weeks)", 
                                        "Chronic (> 1 month)", "Unknown"])
        
        symptom_progression = st.selectbox("Symptom Progression", 
                                         ["Improving", "Stable", "Worsening", "Fluctuating", "Unknown"])
    
    with col2:
        neurological_symptoms = st.multiselect("Neurological Symptoms", 
                                             ["Headache", "Seizures", "Visual Changes", 
                                              "Motor Deficits", "Sensory Changes", 
                                              "Speech Problems", "Memory Loss", "Confusion",
                                              "Nausea/Vomiting", "Balance Issues", "None"])
        
        treatment_history = st.text_area("Treatment History", 
                                       height=100,
                                       placeholder="Previous treatments for current condition...")
        
        additional_notes = st.text_area("Additional Clinical Notes", 
                                      height=100,
                                      placeholder="Any other relevant information...")

# Update the patient_info dictionary to include all the new fields
patient_info = {
    'name': patient_name,
    'age': patient_age,
    'dob': patient_dob,
    'id': patient_id,
    'sex': patient_sex,
    'contact': patient_contact,
    'email': patient_email,
    'address': patient_address,
    'medical_history': medical_history,
    'family_history': family_history,
    'chief_complaint': chief_complaint,
    'current_medications': current_medications,
    'allergies': allergies,
    'previous_imaging': previous_imaging,
    'referring_physician': referring_physician,
    'clinical_diagnosis': clinical_diagnosis,
    'onset_of_symptoms': onset_of_symptoms,
    'symptom_progression': symptom_progression,
    'neurological_symptoms': ', '.join(neurological_symptoms) if neurological_symptoms else "None reported",
    'treatment_history': treatment_history,
    'additional_notes': additional_notes
}
        
def generate_radiology_report(patient_details, scan_findings):
    """
    Function to generate a comprehensive radiological report with Follow-Up and Recommendations sections.
    """
    # Extract relevant information from the input
    patient_name = patient_details.get("name", "Unknown")
    patient_id = patient_details.get("id", "N/A")
    patient_age = patient_details.get("age", "Unknown")
    patient_sex = patient_details.get("sex", "Unknown")
    medical_history = patient_details.get("medical_history", "No history provided")
    
    affected_area = scan_findings.get("affected_percentage", 0)
    location = scan_findings.get("location", "Unknown")
    num_images = scan_findings.get("num_images", 0)
    study_date = datetime.now().strftime('%Y-%m-%d')
    
    # Determine lesion type based on affected area and metrics
    # This is a simplistic approach - in a real application, you would use more sophisticated criteria
    volume = scan_findings.get("volume_estimate", 0)
    diameter = scan_findings.get("max_diameter", 0)
    irregularity = scan_findings.get("shape_irregularity", 0)
    
    # Simple heuristic to determine lesion type
    if affected_area < 2.0 and irregularity < 1.5:
        lesion_type = "Benign"
    elif affected_area > 5.0 or irregularity > 2.5:
        lesion_type = "Malignant"
    else:
        lesion_type = "Indeterminate"
    
    # Define Follow-Up and Recommendation Logic
    def follow_up_recommendation(lesion_type):
        if lesion_type == "Benign":
            return "Follow-up MRI in 6 months to monitor any changes in size or characteristics."
        elif lesion_type == "Malignant":
            return "Immediate referral to oncology for biopsy and further management. Follow-up MRI in 1-2 months post-treatment."
        else:
            return "Further evaluation recommended. Follow-up in 3 months with contrast-enhanced imaging."
    
    # Get the Follow-Up and Recommendation text based on lesion type
    follow_up = follow_up_recommendation(lesion_type)
    
    # Customize recommendations based on patient characteristics and findings
    recommendations = []
    if lesion_type == "Malignant":
        recommendations.append("Neurosurgical consultation is recommended within 1 week.")
        recommendations.append("Consider stereotactic biopsy for histopathological confirmation.")
        recommendations.append("Full neurological assessment is advised.")
    elif lesion_type == "Benign":
        recommendations.append("Clinical monitoring of symptoms is advised.")
        recommendations.append("Consider contrast-enhanced MRI if symptoms worsen.")
    else:
        recommendations.append("Additional imaging studies including contrast-enhanced MRI are recommended.")
        recommendations.append("Clinical correlation with presenting symptoms is advised.")
    
    recommendation_text = "\n- ".join(recommendations)
    
    # Generate the full report with Follow-Up and Recommendations included
    report = f"""RADIOLOGICAL REPORT

CLINICAL HISTORY
{medical_history}

TECHNIQUE
Brain MRI was performed without contrast enhancement. {num_images} image(s) were analyzed.
Date of study: {study_date}

FINDINGS
Brain MRI demonstrates an abnormal area of signal intensity in the {location} of the brain. The lesion affects approximately {affected_area:.2f}% of the brain tissue in the visualized area. The lesion has an estimated volume of {volume} pixels and maximum diameter of {diameter} pixels. The shape irregularity index is {irregularity:.2f}, suggesting a {'irregular' if irregularity > 1.8 else 'relatively regular'} morphology.

No significant mass effect or midline shift is observed. No evidence of hemorrhage or acute infarction.

IMPRESSION
Findings are consistent with a {lesion_type.lower()} space-occupying lesion in the {location}. Differential diagnosis includes:
- {'Primary brain neoplasm' if lesion_type == 'Malignant' else 'Benign meningioma'}
- {'Metastatic disease' if lesion_type == 'Malignant' else 'Low-grade glioma'}
- {'High-grade glioma' if lesion_type == 'Malignant' else 'Focal cortical dysplasia'}

FOLLOW-UP
{follow_up}

RECOMMENDATIONS
- {recommendation_text}

CRITICAL FINDINGS NOTIFICATION
{'Critical findings were communicated to the referring physician.' if lesion_type == 'Malignant' else 'No critical findings requiring immediate notification.'}

Report generated on {datetime.now().strftime('%Y-%m-%d at %H:%M')}
"""
    
    return report    
# In the "Generate Report" section, replace or augment the existing report options
st.subheader("Report Options")

report_type = st.radio(
    "Select Report Type",
    ["AI-Generated Report", "Template-Based Report", "Custom Radiology Report"]
)

if st.button("Generate AI Report"):
    import ollama  # Local model library

    def generate_local_report(prompt, model="llama3.2"):
        try:
            response = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
            return response['message']['content']
        except Exception as e:
            return f"Error: {e}"

    try:
        with st.spinner("Generating AI report..."):
            prompt = f"""
Generate a comprehensive radiological report for a brain scan with the following details:

**Patient Information:**
- **Patient Name:** {patient_info['name']}
- **Patient ID:** {patient_info['id']}
- **Age:** {patient_info['age']}
- **Sex:** {patient_info['sex']}
- **Medical History:** {patient_info['medical_history']}
- **Family History:** {patient_info['family_history']}
- **Chief Complaint/Symptoms:** {patient_info['chief_complaint']}
- **Current Medications:** {patient_info['current_medications']}
- **Known Allergies:** {patient_info['allergies']}
- **Previous Imaging Studies:** {patient_info['previous_imaging']}
- **Referring Physician:** {patient_info['referring_physician']}
- **Provisional Diagnosis:** {patient_info['clinical_diagnosis']}
- **Onset of Symptoms:** {patient_info['onset_of_symptoms']}
- **Symptom Progression:** {patient_info['symptom_progression']}
- **Neurological Symptoms:** {patient_info['neurological_symptoms']}
- **Treatment History:** {patient_info['treatment_history']}
- **Additional Clinical Notes:** {patient_info['additional_notes']}

**Scan Findings:**
- Affected brain area percentage: {st.session_state['segmentation_results']['affected_percentage']:.2f}%
- Location: {st.session_state['segmentation_results']['location']}
- Number of analyzed images: {st.session_state['segmentation_results']['num_images']}
- Date of study: {datetime.now().strftime('%Y-%m-%d')}

Provide:
1. Clinical Summary
2. Technique
3. Findings
4. Impression
5. Recommendations
"""

            ai_report = generate_local_report(prompt, model="llama3.2")

            if ai_report:
                st.session_state['report_text'] = ai_report
                st.success("AI-generated report ready!")
            else:
                st.error("Failed to generate report from local model.")

    except Exception as e:
        st.error(f"Error generating AI report: {e}")

elif report_type == "Custom Radiology Report":
    # Additional options for the custom report
    st.subheader("Lesion Classification")
    lesion_classification = st.selectbox(
        "Lesion Type Assessment",
        ["Automatic", "Benign", "Malignant", "Indeterminate"]
    )
    
    if st.button("Generate Custom Radiology Report"):
        # Prepare patient details and scan findings
        if lesion_classification != "Automatic":
            # Override the automatic classification
            st.session_state['segmentation_results']['lesion_type'] = lesion_classification
        
        # Generate the report
        custom_report = generate_radiology_report(patient_info, st.session_state['segmentation_results'])
        st.session_state['report_text'] = custom_report
        st.success("Custom radiology report generated!")
        # Display report preview
if 'report_text' in st.session_state:
            st.subheader("Report Preview")
            st.text_area("Report Content", st.session_state['report_text'], height=400)
            
            # Export options
            st.subheader("Export Options")
            
            export_format = st.selectbox("Select Export Format", ["PDF", "DOCX", "HTML"])
            
            if st.button("Generate and Download Report"):
                try:
                    if export_format == "PDF":
                        buffer = create_pdf_report(
                            patient_info, 
                            st.session_state['segmentation_results'], 
                            st.session_state['report_text'],
                            st.session_state.get('image_data', None)
                        )
                        st.download_button(
                            label="Download PDF Report",
                            data=buffer,
                            file_name=f"report_{patient_info['id']}_{datetime.now().strftime('%Y%m%d')}.pdf",
                            mime="application/pdf"
                        )
                    elif export_format == "DOCX":
                        buffer = create_docx_report(
                            patient_info, 
                            st.session_state['segmentation_results'], 
                            st.session_state['report_text'],
                            st.session_state.get('image_data', None)
                        )
                        st.download_button(
                            label="Download DOCX Report",
                            data=buffer,
                            file_name=f"report_{patient_info['id']}_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    elif export_format == "HTML":
                        html_content = create_html_report(
                            patient_info, 
                            st.session_state['segmentation_results'], 
                            st.session_state['report_text'],
                            st.session_state.get('image_data', None)
                        )
                        st.download_button(
                            label="Download HTML Report",
                            data=html_content,
                            file_name=f"report_{patient_info['id']}_{datetime.now().strftime('%Y%m%d')}.html",
                            mime="text/html"
                        )
                        
                        # Display HTML preview
                        with st.expander("HTML Preview"):
                            st.components.v1.html(html_content, height=600)
                    
                    st.success(f"Report successfully generated in {export_format} format!")
                except Exception as e:
                    st.error(f"Error generating report: {e}")


# Run the main application
if __name__ == "__main__":
    main()
